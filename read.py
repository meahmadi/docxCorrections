# -*- coding: utf-8 -*-
from docx import *
import re


def paragraph_replace(doc, search, replace):
	searchre = re.compile(search)
	for paragraph in doc.paragraphs:
		paragraph_text = paragraph.text
		if paragraph_text:
			if searchre.search(paragraph_text):
				clear_paragraph(doc,paragraph)
				paragraph.add_run(re.sub(search, replace, paragraph_text))
	return paragraph

def clear_paragraph(doc, paragraph):
	p_element = paragraph._p
	p_child_elements = [elm for elm in p_element.iterchildren()]
	for child_element in p_child_elements:
		p_element.remove(child_element)

doc = Document("yasin.docx")
paragraph_replace(doc,u"1",u"۱")
paragraph_replace(doc,u"2",u"۲")
paragraph_replace(doc,u"3",u"۳")
paragraph_replace(doc,u"4",u"۴")
paragraph_replace(doc,u"5",u"۵")
paragraph_replace(doc,u"6",u"۶")
paragraph_replace(doc,u"7",u"۷")
paragraph_replace(doc,u"8",u"۸")
paragraph_replace(doc,u"9",u"۹")
paragraph_replace(doc,u"0",u"۰")

doc.save("yasin1.docx")